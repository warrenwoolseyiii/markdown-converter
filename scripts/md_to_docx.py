#!/usr/bin/env python3
"""
Markdown to Word Document Converter

Converts Markdown files to Microsoft Word (.docx) format,
using a template document to preserve all styles, headers, fonts, etc.

Usage:
    python md_to_docx.py input.md output.docx --template template.docx
"""

import argparse
import re
import sys
import copy
import yaml
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class MarkdownParser:
    """Parse Markdown content into structured elements."""
    
    def __init__(self, content: str):
        self.content = content
        self.elements = []
        self.front_matter = {}
        
    def parse(self) -> List[Dict[str, Any]]:
        """Parse the markdown content into a list of elements."""
        # First, extract YAML front matter if present
        content = self._extract_front_matter(self.content)
        
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # Check for code blocks (```)
            if line.strip().startswith('```'):
                code_content, end_idx = self._parse_code_block(lines, i)
                self.elements.append({
                    'type': 'code_block',
                    'content': code_content
                })
                i = end_idx + 1
                continue
            
            # Check for tables
            if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
                table_data, end_idx = self._parse_table(lines, i)
                if table_data:
                    self.elements.append({
                        'type': 'table',
                        'content': table_data
                    })
                    i = end_idx + 1
                    continue
            
            # Check for headings
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                self.elements.append({
                    'type': 'heading',
                    'level': level,
                    'content': text
                })
                i += 1
                continue
            
            # Check for page break comment
            if re.match(r'^\s*<!--\s*pagebreak\s*-->\s*$', line.strip(), re.IGNORECASE):
                self.elements.append({
                    'type': 'page_break'
                })
                i += 1
                continue
            
            # Check for TOC placeholder comment
            if re.match(r'^\s*<!--\s*toc\s*-->\s*$', line.strip(), re.IGNORECASE):
                self.elements.append({
                    'type': 'toc'
                })
                i += 1
                continue
            
            # Check for horizontal rule
            if re.match(r'^[-*_]{3,}\s*$', line.strip()):
                self.elements.append({
                    'type': 'horizontal_rule'
                })
                i += 1
                continue
            
            # Check for numbered list
            numbered_match = re.match(r'^(\d+)\.\s+(.+)$', line)
            if numbered_match:
                items, end_idx = self._parse_numbered_list(lines, i)
                self.elements.append({
                    'type': 'numbered_list',
                    'content': items
                })
                i = end_idx + 1
                continue
            
            # Check for bullet list
            bullet_match = re.match(r'^[-*+]\s+(.+)$', line)
            if bullet_match:
                items, end_idx = self._parse_bullet_list(lines, i)
                self.elements.append({
                    'type': 'bullet_list',
                    'content': items
                })
                i = end_idx + 1
                continue
            
            # Regular paragraph
            if line.strip():
                para_content, end_idx = self._parse_paragraph(lines, i)
                self.elements.append({
                    'type': 'paragraph',
                    'content': para_content
                })
                i = end_idx + 1
                continue
            
            # Empty line
            i += 1
        
        return self.elements
    
    def _parse_code_block(self, lines: List[str], start: int) -> Tuple[str, int]:
        """Parse a code block starting at the given index."""
        code_lines = []
        i = start + 1  # Skip the opening ```
        
        while i < len(lines):
            if lines[i].strip().startswith('```'):
                break
            code_lines.append(lines[i])
            i += 1
        
        return '\n'.join(code_lines), i
    
    def _parse_table(self, lines: List[str], start: int) -> Tuple[List[List[str]], int]:
        """Parse a markdown table starting at the given index."""
        table_rows = []
        i = start
        
        while i < len(lines) and '|' in lines[i]:
            line = lines[i].strip()
            
            # Skip separator line (|---|---|)
            if re.match(r'^[\|\s\-:]+$', line):
                i += 1
                continue
            
            # Parse cells
            cells = [cell.strip() for cell in line.split('|')]
            # Remove empty first and last cells from leading/trailing |
            if cells and cells[0] == '':
                cells = cells[1:]
            if cells and cells[-1] == '':
                cells = cells[:-1]
            
            if cells:
                table_rows.append(cells)
            i += 1
        
        return table_rows, i - 1
    
    def _parse_numbered_list(self, lines: List[str], start: int) -> Tuple[List[Dict], int]:
        """Parse a numbered list starting at the given index.
        
        Returns a list of dicts with 'text' and optional 'subitems' (nested bullets).
        """
        items = []
        i = start
        
        while i < len(lines):
            line = lines[i]
            
            # Check for numbered list item
            match = re.match(r'^(\d+)\.\s+(.+)$', line)
            if match:
                item = {'text': match.group(2), 'subitems': []}
                items.append(item)
                i += 1
                continue
            
            # Check for nested bullet under last numbered item (indented with spaces)
            nested_bullet = re.match(r'^(\s{2,})[-*+]\s+(.+)$', line)
            if nested_bullet and items:
                # Add as subitem to last numbered item
                items[-1]['subitems'].append(nested_bullet.group(2))
                i += 1
                continue
            
            # Check for continuation of previous item (indented text, not a bullet)
            if line.startswith('   ') and items and not re.match(r'^\s*[-*+]\s', line):
                items[-1]['text'] += ' ' + line.strip()
                i += 1
                continue
            
            # Empty line - might be separator between items
            if not line.strip():
                # Look ahead to see if next line is another numbered item
                if i + 1 < len(lines) and re.match(r'^(\d+)\.\s+', lines[i + 1]):
                    i += 1
                    continue
                else:
                    break
            
            # Not part of this list
            break
        
        return items, i - 1
    
    def _parse_bullet_list(self, lines: List[str], start: int) -> Tuple[List[Dict], int]:
        """Parse a bullet list starting at the given index.
        
        Returns a list of dicts with 'text' and optional 'subitems' (nested bullets).
        Handles nested bullets like:
        - Main item
          - Sub item 1
          - Sub item 2
        """
        items = []
        i = start
        
        while i < len(lines):
            line = lines[i]
            
            # Check for main bullet item (not indented or minimal indent)
            main_bullet = re.match(r'^[-*+]\s+(.+)$', line)
            if main_bullet:
                item = {'text': main_bullet.group(1), 'subitems': []}
                items.append(item)
                i += 1
                continue
            
            # Check for nested bullet (indented with 2+ spaces before the bullet)
            nested_bullet = re.match(r'^(\s{2,})[-*+]\s+(.+)$', line)
            if nested_bullet and items:
                # Add as subitem to last bullet item
                items[-1]['subitems'].append(nested_bullet.group(2))
                i += 1
                continue
            
            # Check for continuation of previous item (indented text, not a bullet)
            if line.startswith('   ') and items and not re.match(r'^\s*[-*+]\s', line):
                items[-1]['text'] += ' ' + line.strip()
                i += 1
                continue
            
            # Empty line - might be separator between items
            if not line.strip():
                # Look ahead to see if next line is another bullet item
                if i + 1 < len(lines) and re.match(r'^[-*+]\s+', lines[i + 1]):
                    i += 1
                    continue
                else:
                    break
            
            # Not part of this list
            break
        
        return items, i - 1
    
    def _parse_paragraph(self, lines: List[str], start: int) -> Tuple[str, int]:
        """Parse a paragraph starting at the given index."""
        para_lines = []
        i = start
        
        while i < len(lines):
            line = lines[i]
            # Stop at empty line or start of another element
            if not line.strip():
                break
            if re.match(r'^#{1,6}\s', line):
                break
            if re.match(r'^[-*_]{3,}\s*$', line):
                break
            if line.strip().startswith('```'):
                break
            if re.match(r'^\d+\.\s', line):
                break
            if re.match(r'^[-*+]\s', line):
                break
            if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
                break
            
            para_lines.append(line)
            i += 1
        
        return ' '.join(para_lines), i - 1
    
    def _extract_front_matter(self, content: str) -> str:
        """Extract YAML front matter from the content if present.
        
        Front matter is delimited by --- at the start and end.
        Returns the content with front matter removed.
        """
        lines = content.split('\n')
        
        if not lines or lines[0].strip() != '---':
            return content
        
        # Find the closing ---
        end_index = -1
        for i in range(1, len(lines)):
            if lines[i].strip() == '---':
                end_index = i
                break
        
        if end_index == -1:
            return content
        
        # Parse YAML front matter
        yaml_content = '\n'.join(lines[1:end_index])
        try:
            self.front_matter = yaml.safe_load(yaml_content) or {}
        except yaml.YAMLError:
            self.front_matter = {}
        
        # Return content without front matter
        return '\n'.join(lines[end_index + 1:])


class MarkdownToDocxConverter:
    """Convert Markdown to Word Document using a template."""
    
    def __init__(self, template_path: Optional[str] = None, auto_number_headings: bool = False,
                 strip_heading_numbers: bool = True, h1_is_title: bool = True,
                 add_title_page: bool = True, add_toc: bool = True, add_approval_section: bool = True):
        self.template_path = template_path
        self.document = None
        self.heading_counters = [0, 0, 0, 0, 0, 0]  # For numbered headings
        self.auto_number_headings = auto_number_headings  # Set to False if markdown already has numbers
        self.strip_heading_numbers = strip_heading_numbers  # Strip numbers from markdown headings (use when template has auto-numbering)
        self.h1_is_title = h1_is_title  # Treat H1 as document title (use Title style, not Heading 1)
        self.add_title_page = add_title_page  # Add a title page from front matter
        self.add_toc = add_toc  # Add table of contents
        self.add_approval_section = add_approval_section  # Add approval section before revision history
        self.front_matter = {}  # Will be populated from parser
        
    def convert(self, markdown_content: str) -> Document:
        """Convert markdown content to a Word document."""
        # Load template or create new document
        if self.template_path and Path(self.template_path).exists():
            print(f"Loading template: {self.template_path}")
            self.document = Document(self.template_path)
            # Clear the template content but keep styles
            self._clear_document_content()
        else:
            print("Creating new document with default styles")
            self.document = Document()
            self._setup_document_defaults()
        
        # Parse markdown
        parser = MarkdownParser(markdown_content)
        elements = parser.parse()
        self.front_matter = parser.front_matter
        
        # Check if markdown contains a TOC placeholder
        has_toc_placeholder = any(e.get('type') == 'toc' for e in elements)
        
        # Add title page if front matter is present and option is enabled
        if self.add_title_page and self.front_matter:
            self._add_title_page()
        
        # Add table of contents if enabled AND no explicit placeholder exists
        # (If placeholder exists, TOC will be added when we process that element)
        if self.add_toc and not has_toc_placeholder:
            self._add_table_of_contents()
        
        # Add approval section if enabled (before main content)
        # Only add if no TOC placeholder (approval goes after TOC)
        if self.add_approval_section and self.front_matter and not has_toc_placeholder:
            self._add_approval_section()
        
        # Track if we've added the approval section after TOC placeholder
        self._approval_added = not self.add_approval_section or not has_toc_placeholder
        
        # Process each element
        for element in elements:
            self._process_element(element)
        
        return self.document
    
    def _clear_document_content(self):
        """Clear the content of the template but preserve styles."""
        # Remove all paragraphs and tables from the body
        body = self.document._body._body
        # Get all p (paragraph) and tbl (table) elements
        for element in body.findall('.//w:p', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            element.getparent().remove(element)
        for element in body.findall('.//w:tbl', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            element.getparent().remove(element)
    
    def _setup_document_defaults(self):
        """Set up default document settings when no template is used."""
        # Set page margins
        for section in self.document.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def _add_title_page(self):
        """Add a title page based on front matter metadata.
        
        Creates a cover page with:
        - Product name (large, centered)
        - Document type (e.g., "Product Specification")
        - Revision number
        """
        fm = self.front_matter
        
        # Add several blank lines for spacing at top
        for _ in range(8):
            para = self.document.add_paragraph()
            para.paragraph_format.space_after = Pt(0)
        
        # Full title
        full_title = fm.get('full_title', fm.get('product_name', 'Product'))
        
        # Add full title - large centered text
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(full_title)
        run.font.size = Pt(24)
        run.font.bold = True
        para.paragraph_format.space_after = Pt(48)
        
        # Add more spacing
        for _ in range(4):
            para = self.document.add_paragraph()
            para.paragraph_format.space_after = Pt(0)
        
        # Document type (e.g., "Product Specification")
        doc_type = fm.get('document_type', 'Specification')
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(doc_type)
        run.font.size = Pt(18)
        para.paragraph_format.space_after = Pt(24)
        
        # Revision number
        revision = fm.get('revision', '0.1')
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"Revision  {revision}")
        run.font.size = Pt(14)
        para.paragraph_format.space_after = Pt(24)
        
        # Add page break after title page
        para = self.document.add_paragraph()
        run = para.add_run()
        run.add_break(WD_BREAK.PAGE)
    
    def _add_table_of_contents(self):
        """Add a Table of Contents field to the document.
        
        This creates a TOC field that Word will populate when the document is opened.
        The user will need to right-click and "Update Field" to refresh the TOC.
        """
        # Add TOC heading
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run("Table of Contents")
        run.font.size = Pt(16)
        run.font.bold = True
        para.paragraph_format.space_after = Pt(12)
        
        # Add TOC field
        para = self.document.add_paragraph()
        
        # Create the TOC field using XML
        # This creates a field that Word will recognize and populate
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '  # TOC with headings 1-3, hyperlinks
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        # Placeholder text - will be replaced when TOC is updated
        placeholder = OxmlElement('w:t')
        placeholder.text = "Right-click and select 'Update Field' to generate Table of Contents"
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        # Build the field in the paragraph's run
        run = para.add_run()
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(placeholder)
        run._r.append(fldChar3)
        
        # Add page break after TOC
        para = self.document.add_paragraph()
        run = para.add_run()
        run.add_break(WD_BREAK.PAGE)
    
    def _add_approval_section(self):
        """Add a 'Reviewed and Approved By' section with signature lines."""
        fm = self.front_matter
        approvers = fm.get('approvers', [])
        
        if not approvers:
            return
        
        # Add heading
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("Reviewed and Approved By")
        run.font.size = Pt(14)
        run.font.bold = True
        para.paragraph_format.space_after = Pt(24)
        
        # Add signature lines in a table (2 columns for approvers)
        num_approvers = len(approvers)
        num_rows = (num_approvers + 1) // 2  # 2 approvers per row
        
        if num_approvers > 0:
            table = self.document.add_table(rows=num_rows, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Set table width
            for row in table.rows:
                for cell in row.cells:
                    cell.width = Inches(3)
            
            # Fill in approver names with signature lines
            for i, approver in enumerate(approvers):
                row_idx = i // 2
                col_idx = i % 2
                
                cell = table.rows[row_idx].cells[col_idx]
                
                # Add blank line for signature
                para = cell.paragraphs[0]
                para.add_run("\n\n")
                
                # Add signature line
                para = cell.add_paragraph()
                para.add_run("_" * 30)
                para.paragraph_format.space_after = Pt(6)
                
                # Add approver name
                para = cell.add_paragraph()
                para.add_run(approver)
                para.paragraph_format.space_after = Pt(24)
        
        # Add spacing after approval section
        para = self.document.add_paragraph()
        para.paragraph_format.space_after = Pt(24)
        
        # Add page break after approval section
        para = self.document.add_paragraph()
        run = para.add_run()
        run.add_break(WD_BREAK.PAGE)
    
    def _get_style_name(self, base_name: str) -> str:
        """Get the actual style name from the document, handling variations."""
        # Try common variations
        variations = [
            base_name,
            base_name.replace(' ', ''),
            base_name.title(),
            base_name.lower(),
        ]
        
        for style_name in variations:
            try:
                style = self.document.styles[style_name]
                return style_name
            except KeyError:
                continue
        
        return base_name  # Return original if none found
    
    def _process_element(self, element: Dict[str, Any]):
        """Process a single parsed element."""
        elem_type = element['type']
        
        if elem_type == 'heading':
            self._add_heading(element['content'], element['level'])
        elif elem_type == 'paragraph':
            self._add_paragraph(element['content'])
        elif elem_type == 'code_block':
            self._add_code_block(element['content'])
        elif elem_type == 'table':
            self._add_table(element['content'])
        elif elem_type == 'numbered_list':
            self._add_numbered_list(element['content'])
        elif elem_type == 'bullet_list':
            self._add_bullet_list(element['content'])
        elif elem_type == 'horizontal_rule':
            self._add_horizontal_rule()
        elif elem_type == 'page_break':
            self._add_page_break()
        elif elem_type == 'toc':
            self._add_toc_from_placeholder()
    
    def _add_heading(self, text: str, level: int):
        """Add a heading to the document using template styles.
        
        When h1_is_title is True:
        - H1 (#) becomes document Title (not numbered)
        - H2 (##) becomes Heading 1 (numbered as 1, 2, 3...)
        - H3 (###) becomes Heading 2 (numbered as 1.1, 1.2, 2.1...)
        - H4 (####) becomes Heading 3 (numbered as 1.1.1, 1.1.2...)
        """
        # Check if heading already starts with a number (e.g., "1. Document Overview" or "1.1 Purpose")
        number_match = re.match(r'^([\d]+[.\d]*\.?\s+)(.+)$', text)
        
        if self.strip_heading_numbers and number_match:
            # Strip the leading numbers from the heading (template will provide auto-numbering)
            full_text = number_match.group(2)
        elif self.auto_number_headings and not number_match:
            # Update heading counters for numbered headings
            self.heading_counters[level - 1] += 1
            for i in range(level, len(self.heading_counters)):
                self.heading_counters[i] = 0
            
            # Build numbered heading
            number_parts = [str(self.heading_counters[i]) for i in range(level) if self.heading_counters[i] > 0]
            heading_number = '.'.join(number_parts)
            
            # Create the heading text with number
            full_text = f"{heading_number}\t{text}"
        else:
            # Use the text as-is
            full_text = text
        
        # Determine the Word style to use
        if self.h1_is_title and level == 1:
            # H1 is treated as a document title - use Title style or just large bold text
            style_name = self._get_style_name('Title')
            try:
                para = self.document.add_paragraph(full_text, style=style_name)
            except KeyError:
                # Fall back to manual title formatting
                para = self.document.add_paragraph()
                run = para.add_run(full_text)
                run.font.bold = True
                run.font.size = Pt(24)
                para.paragraph_format.space_after = Pt(12)
        else:
            # Shift heading level down by 1 if h1_is_title is enabled
            # H2 -> Heading 1, H3 -> Heading 2, H4 -> Heading 3, etc.
            word_level = level - 1 if self.h1_is_title else level
            word_level = max(1, word_level)  # Ensure at least level 1
            
            style_name = self._get_style_name(f'Heading {word_level}')
            
            try:
                para = self.document.add_paragraph(full_text, style=style_name)
            except KeyError:
                # Fall back to manual formatting if style not found
                para = self.document.add_paragraph(full_text)
                run = para.runs[0] if para.runs else para.add_run()
                run.font.bold = True
                run.font.size = Pt(16 - (word_level - 1) * 2)  # Decreasing size for deeper levels
    
    def _add_paragraph(self, text: str):
        """Add a paragraph with inline formatting."""
        para = self.document.add_paragraph()
        
        # Try to apply Normal style from template
        try:
            para.style = self.document.styles['Normal']
        except KeyError:
            pass
        
        self._add_formatted_text(para, text)
    
    def _add_hyperlink(self, paragraph, text: str, url: str):
        """Add a hyperlink to a paragraph.
        
        Creates a proper Word hyperlink with blue underlined text.
        """
        # Get the document's relationships part
        part = self.document.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        # Create the hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        # Create a new run for the hyperlink text
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # Set hyperlink styling (blue color, underline)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')  # Blue
        rPr.append(color)
        
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        rPr.append(underline)
        
        new_run.append(rPr)
        
        # Add the text
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)
        
        hyperlink.append(new_run)
        
        # Append the hyperlink to the paragraph
        paragraph._p.append(hyperlink)
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text with inline formatting (bold, italic, code, hyperlinks) to a paragraph."""
        # Pattern for:
        # - Links: [text](url)
        # - Bold: **text**
        # - Italic: *text* or _text_
        # - Code: `text`
        # Note: Link pattern must come first to avoid conflict with other patterns
        link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
        format_pattern = r'(\*\*(.+?)\*\*)|(\*(.+?)\*)|(_(.+?)_)|(`(.+?)`)'
        
        # Combined pattern - links first, then formatting
        combined_pattern = f'({link_pattern})|{format_pattern}'
        
        last_end = 0
        for match in re.finditer(combined_pattern, text):
            # Add text before the match
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])
            
            # Check if it's a hyperlink match
            if match.group(2) and match.group(3):  # Link [text](url)
                link_text = match.group(2)
                link_url = match.group(3)
                self._add_hyperlink(paragraph, link_text, link_url)
            # Determine formatting type
            elif match.group(5):  # Bold (**text**)
                run = paragraph.add_run(match.group(5))
                run.font.bold = True
            elif match.group(7):  # Italic (*text*)
                run = paragraph.add_run(match.group(7))
                run.font.italic = True
            elif match.group(9):  # Italic (_text_)
                run = paragraph.add_run(match.group(9))
                run.font.italic = True
            elif match.group(11):  # Code (`text`)
                run = paragraph.add_run(match.group(11))
                run.font.name = 'Consolas'
            
            last_end = match.end()
        
        # Add remaining text
        if last_end < len(text):
            paragraph.add_run(text[last_end:])
    
    def _add_code_block(self, content: str):
        """Add a code block (e.g., ASCII diagrams) to the document.
        
        Uses a smaller font size for wide ASCII diagrams to fit within page margins.
        Standard Word page with 1" margins = ~6.5" usable width.
        At 7pt Consolas, approximately 95-100 characters fit per line.
        At 8pt Consolas, approximately 85-90 characters fit per line.
        """
        # Determine the maximum line width to select appropriate font size
        lines = content.split('\n')
        max_line_length = max(len(line) for line in lines) if lines else 0
        
        # Select font size based on content width
        # Standard page width allows ~65a chars at 9pt, ~85 chars at 7pt with Consolas
        if max_line_length > 90:
            font_size = Pt(6.5)  # Very wide diagrams - smallest readable size
        elif max_line_length > 75:
            font_size = Pt(7)    # Wide diagrams
        elif max_line_length > 65:
            font_size = Pt(8)    # Medium-wide diagrams
        else:
            font_size = Pt(9)    # Normal code blocks
        
        # Add each line as a separate paragraph to preserve formatting
        for line in lines:
            para = self.document.add_paragraph()
            run = para.add_run(line if line else ' ')  # Use space for empty lines
            
            # Apply monospace font for code with calculated size
            run.font.name = 'Consolas'
            run.font.size = font_size
            
            # Set tight line spacing for code
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
    
    def _add_table(self, data: List[List[str]]):
        """Add a table to the document."""
        if not data:
            return
        
        # Create table
        rows = len(data)
        cols = max(len(row) for row in data) if data else 0
        
        if cols == 0:
            return
        
        table = self.document.add_table(rows=rows, cols=cols)
        
        # Try to use a table style from the template
        for style_name in ['Table Grid', 'Grid Table 1 Light', 'Table Normal']:
            try:
                table.style = style_name
                break
            except KeyError:
                continue
        
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Fill table data
        for i, row_data in enumerate(data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = ''
                    para = cell.paragraphs[0]
                    
                    # Add formatted text
                    self._add_formatted_text(para, cell_text)
                    
                    # Style header row differently
                    if i == 0:
                        for run in para.runs:
                            run.font.bold = True
                        self._set_cell_shading(cell, 'D3D3D3')  # Light gray
        
        # Add spacing after table
        para = self.document.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
    
    def _set_cell_shading(self, cell, color_hex: str):
        """Set background color for a table cell."""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex)
        cell._tc.get_or_add_tcPr().append(shading)
    
    def _create_numbering_part(self):
        """Create or get the numbering part for the document and set up list definitions."""
        from docx.oxml.ns import nsmap
        
        # Get or create numbering part
        numbering_part = self.document.part.numbering_part
        
        # We'll use the existing numbering definitions if available
        # The numbering IDs will be tracked for creating proper lists
        return numbering_part
    
    def _set_paragraph_numbering(self, para, num_id: int, ilvl: int = 0):
        """Apply numbering to a paragraph using XML manipulation.
        
        Args:
            para: The paragraph to apply numbering to
            num_id: The numbering definition ID
            ilvl: The indentation level (0 for main items, 1 for subitems, etc.)
        """
        # Create the numPr element
        numPr = OxmlElement('w:numPr')
        
        # Set the indentation level
        ilvlElem = OxmlElement('w:ilvl')
        ilvlElem.set(qn('w:val'), str(ilvl))
        numPr.append(ilvlElem)
        
        # Set the numbering ID
        numIdElem = OxmlElement('w:numId')
        numIdElem.set(qn('w:val'), str(num_id))
        numPr.append(numIdElem)
        
        # Add to paragraph properties
        pPr = para._p.get_or_add_pPr()
        pPr.insert(0, numPr)
    
    def _add_numbered_list(self, items: List[Dict]):
        """Add a numbered list to the document.
        
        Items can be either strings or dicts with 'text' and optional 'subitems'.
        Uses Word's built-in numbering with proper indentation.
        """
        # Try to use existing list style, or use Body Text Indent with numbering
        try:
            body_indent_style = self.document.styles['Body Text Indent']
        except KeyError:
            body_indent_style = None
        
        # Use numbering ID 1 for main numbered list (Word's default numbered list)
        # Most documents have this as a decimal numbered list
        num_id = 1
        
        for i, item in enumerate(items, 1):
            # Handle both old (string) and new (dict) formats
            if isinstance(item, str):
                item_text = item
                subitems = []
            else:
                item_text = item.get('text', '')
                subitems = item.get('subitems', [])
            
            # Create paragraph with proper indentation
            para = self.document.add_paragraph()
            
            # Apply Body Text Indent style if available for consistent formatting
            if body_indent_style:
                para.style = body_indent_style
            
            # Apply numbering via XML
            try:
                self._set_paragraph_numbering(para, num_id, ilvl=0)
            except Exception:
                # Fallback: add number manually if numbering fails
                para.add_run(f"{i}.\t")
            
            self._add_formatted_text(para, item_text)
            
            # Add subitems with lower-case letter numbering (a, b, c)
            for j, subitem in enumerate(subitems):
                sub_para = self.document.add_paragraph()
                
                # Apply indented style
                try:
                    sub_para.style = self.document.styles.get('Body Text Indent 2', 'Normal')
                except:
                    pass
                
                # Apply nested numbering (ilvl=1 for second level - typically letters)
                try:
                    self._set_paragraph_numbering(sub_para, num_id, ilvl=1)
                except Exception:
                    # Fallback: add letter manually
                    letter = chr(ord('a') + j)
                    sub_para.add_run(f"    {letter}.\t")
                
                self._add_formatted_text(sub_para, subitem)
    
    def _add_bullet_list(self, items: List[Dict]):
        """Add a bullet list to the document.
        
        Items can be either strings or dicts with 'text' and optional 'subitems'.
        Uses manual bullet formatting with proper indentation to avoid numbered list issues.
        """
        for item in items:
            # Handle both old (string) and new (dict) formats
            if isinstance(item, str):
                item_text = item
                subitems = []
            else:
                item_text = item.get('text', '')
                subitems = item.get('subitems', [])
            
            # Create paragraph for main bullet item
            para = self.document.add_paragraph()
            
            # Use manual bullet character with proper indentation
            # This avoids issues with template styles that may apply numbered lists
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            run = para.add_run("•\t")
            run.font.bold = False  # Ensure bullet is not bold
            
            self._add_formatted_text(para, item_text)
            
            # Add subitems as nested bullets with further indentation
            for subitem in subitems:
                sub_para = self.document.add_paragraph()
                
                # Use manual nested bullet with deeper indentation
                sub_para.paragraph_format.left_indent = Inches(0.75)
                sub_para.paragraph_format.first_line_indent = Inches(-0.25)
                run = sub_para.add_run("○\t")  # Use different bullet for nested items
                run.font.bold = False
                
                self._add_formatted_text(sub_para, subitem)
    
    def _add_horizontal_rule(self):
        """Add a horizontal rule/line to the document."""
        para = self.document.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)
        
        # Add a bottom border to simulate horizontal rule
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        para._p.get_or_add_pPr().append(pBdr)
    
    def _add_page_break(self):
        """Add a page break to the document.
        
        Creates an explicit page break that forces content to start on a new page.
        This is triggered by the <!-- pagebreak --> comment in markdown.
        """
        para = self.document.add_paragraph()
        run = para.add_run()
        run.add_break(WD_BREAK.PAGE)
    
    def _add_toc_from_placeholder(self):
        """Add a Table of Contents at the position of the <!-- toc --> placeholder.
        
        This method is called when processing a TOC placeholder in the markdown.
        It ensures the TOC appears at the correct position in the document flow.
        """
        # Add the table of contents
        if self.add_toc:
            self._add_table_of_contents()
        
        # Add approval section after TOC if not already added
        if self.add_approval_section and self.front_matter and not getattr(self, '_approval_added', True):
            self._add_approval_section()
            self._approval_added = True
    
    def save(self, output_path: str):
        """Save the document to a file."""
        if self.document:
            self.document.save(output_path)
            print(f"Document saved to: {output_path}")


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description='Convert Markdown to Word Document',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
    python md_to_docx.py input.md output.docx
    python md_to_docx.py input.md output.docx --template example_format.docx
        '''
    )
    
    parser.add_argument('input', help='Input Markdown file path')
    parser.add_argument('output', help='Output Word document path')
    parser.add_argument(
        '--template', '-t',
        help='Template document to use for styles (header, fonts, etc.)',
        default=None
    )
    
    args = parser.parse_args()
    
    # Validate input file
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: Input file not found: {args.input}")
        sys.exit(1)
    
    # Create output directory if needed
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Read markdown content
    print(f"Reading: {args.input}")
    with open(input_path, 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    # Convert
    print(f"Converting with template: {args.template or 'default styles'}")
    converter = MarkdownToDocxConverter(template_path=args.template)
    converter.convert(markdown_content)
    
    # Save
    converter.save(args.output)
    print("Conversion complete!")


if __name__ == '__main__':
    main()
